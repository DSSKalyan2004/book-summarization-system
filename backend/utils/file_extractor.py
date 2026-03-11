import PyPDF2
from docx import Document
import re
from pathlib import Path


def _smart_sentence_join(text: str) -> str:
    """Join broken lines into proper sentences with correct punctuation."""
    lines = text.split("\n")
    result_parts = []
    current = ""

    for line in lines:
        stripped = line.strip()
        if not stripped:
            if current:
                result_parts.append(current)
                current = ""
            continue
        if current:
            if current[-1] in ".!?:":
                result_parts.append(current)
                current = stripped
            elif current[-1] == "-":
                current = current[:-1] + stripped
            else:
                current = current + " " + stripped
        else:
            current = stripped

    if current:
        result_parts.append(current)

    return "\n\n".join(result_parts)


def clean_extracted_text(text: str) -> str:
    """Clean extracted text — preserve meaningful punctuation and structure."""
    cleaned = text

    # Remove URLs and emails
    cleaned = re.sub(r'https?://[^\s]+', '', cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r'www\.[^\s]+', '', cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r'[\w.-]+@[\w.-]+\.\w+', '', cleaned, flags=re.IGNORECASE)

    # Remove reference markers [1] (2) etc.
    cleaned = re.sub(r'\[\d+\]', '', cleaned)
    cleaned = re.sub(r'\(\d+\)', '', cleaned)

    # Remove special chars but KEEP proper punctuation
    cleaned = re.sub(r'[^\w\s.,!?;:()\-\'"—–\n]', ' ', cleaned)

    # Remove bullet symbols
    cleaned = re.sub(r'^[\s]*[•▪▫■□●○◆◇★☆►▸]+[\s]*', '', cleaned, flags=re.MULTILINE)
    cleaned = re.sub(r'^[\s]*[\u2022\u2023\u25E6\u2043\u2219]+[\s]*', '', cleaned, flags=re.MULTILINE)

    # Remove standalone page numbers
    cleaned = re.sub(r'^\s*\d{1,4}\s*$', '', cleaned, flags=re.MULTILINE)

    # Fix excessive dots
    cleaned = re.sub(r'\.{4,}', '...', cleaned)
    cleaned = re.sub(r'\.{2}(?!\.)', '.', cleaned)

    # Fix punctuation spacing
    cleaned = re.sub(r'\s+([.,!?;:])', r'\1', cleaned)
    cleaned = re.sub(r'([.,!?;:])([A-Za-z])', r'\1 \2', cleaned)

    # Fix multiple spaces
    cleaned = re.sub(r'[ \t]+', ' ', cleaned)

    # Trim lines
    cleaned = re.sub(r'^[ \t]+|[ \t]+$', '', cleaned, flags=re.MULTILINE)

    # Collapse 3+ newlines into 2
    cleaned = re.sub(r'\n{3,}', '\n\n', cleaned)

    # Smart sentence joining
    cleaned = _smart_sentence_join(cleaned)

    # Capitalize after sentence-ending punctuation
    def _fix_cap(m):
        return m.group(1) + " " + m.group(2).upper()
    cleaned = re.sub(r'([.!?])\s+([a-z])', _fix_cap, cleaned)

    # Fix common OCR artefacts
    cleaned = re.sub(r'\bﬁ', 'fi', cleaned)
    cleaned = re.sub(r'\bﬂ', 'fl', cleaned)
    cleaned = re.sub(r'\bﬀ', 'ff', cleaned)
    cleaned = re.sub(r'[\u201c\u201d]', '"', cleaned)
    cleaned = re.sub(r'[\u2018\u2019]', "'", cleaned)
    cleaned = re.sub(r'[–—]', ' - ', cleaned)
    cleaned = re.sub(r'  +', ' ', cleaned)

    return cleaned.strip()


async def extract_text_from_file(file_path: str, file_extension: str) -> str:
    """Extract text from uploaded files (PDF, DOCX, TXT) with high quality."""
    ext = file_extension.lower().replace('.', '')

    if ext == 'pdf':
        raw_text = await extract_from_pdf(file_path)
    elif ext == 'docx':
        raw_text = await extract_from_docx(file_path)
    elif ext == 'txt':
        raw_text = await extract_from_txt(file_path)
    else:
        raise ValueError(f"Unsupported file format: {file_extension}. Supported: PDF, DOCX, TXT")

    cleaned_text = clean_extracted_text(raw_text)
    print(f"🧹 Cleaned text: {len(raw_text)} → {len(cleaned_text)} characters")

    return cleaned_text


async def extract_from_pdf(file_path: str) -> str:
    """Extract text from PDF with improved paragraph detection."""
    try:
        text_parts = []
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)

            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if not page_text:
                    continue
                # Merge hyphenated line breaks
                page_text = re.sub(r'-\s*\n\s*', '', page_text)
                # Replace single newlines (within paragraph) with space
                page_text = re.sub(r'(?<!\n)\n(?!\n)', ' ', page_text)
                # Normalize whitespace within lines
                page_text = re.sub(r'[ \t]+', ' ', page_text)
                text_parts.append(page_text.strip())

        text = "\n\n".join(text_parts)

        if not text or text.strip() == "":
            raise ValueError("No text content found in PDF file")

        print(f"✅ Extracted {len(text)} characters from PDF ({len(pdf_reader.pages)} pages)")
        return text
    except ValueError:
        raise
    except Exception as e:
        raise ValueError(f"Failed to extract text from PDF: {str(e)}")


async def extract_from_docx(file_path: str) -> str:
    """Extract text from DOCX with table and heading support."""
    try:
        doc = Document(file_path)
        parts = []

        for para in doc.paragraphs:
            txt = para.text.strip()
            if not txt:
                continue
            if para.style and para.style.name and para.style.name.startswith("Heading"):
                parts.append("")
            parts.append(txt)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

        text = "\n".join(parts)

        if not text or text.strip() == "":
            raise ValueError("No text content found in DOCX file")

        print(f"✅ Extracted {len(text)} characters from DOCX")
        return text
    except ValueError:
        raise
    except Exception as e:
        raise ValueError(f"Failed to extract text from DOCX: {str(e)}")


async def extract_from_txt(file_path: str) -> str:
    """Extract text from TXT files with encoding fallback."""
    try:
        for encoding in ('utf-8', 'utf-8-sig', 'latin-1', 'cp1252'):
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    text = file.read()
                if text and text.strip():
                    print(f"✅ Extracted {len(text)} characters from TXT ({encoding})")
                    return text
            except (UnicodeDecodeError, UnicodeError):
                continue

        raise ValueError("No text content found in TXT file")
    except ValueError:
        raise
    except Exception as e:
        raise ValueError(f"Failed to extract text from TXT: {str(e)}")
